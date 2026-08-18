import streamlit as st
import subprocess
import re
import shutil
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# --- 1. Page Config & Custom CSS ---
st.set_page_config(page_title="Solid Black | Paper Generator", layout="wide", page_icon="📝")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Hind+Vadodara:wght@400;600;700&display=swap');
    html, body, [class*="css"], p, h1, h2, h3, h4, span, label { font-family: 'Hind Vadodara', sans-serif !important; }
    .main-title { text-align: center; font-weight: 700; font-size: 32px; margin-top: 5px; margin-bottom: 5px; color: #1a1a1a; }
    .subtitle-badge { text-align: center; margin-bottom: 25px; }
    .subtitle-badge span { background-color: #1a1a1a; color: #ffffff; padding: 5px 15px; border-radius: 20px; font-size: 13px; letter-spacing: 1px; }
    div.stButton > button:first-child { background-color: #1F4E79; color: #ffffff; border-radius: 6px; padding: 10px 24px; font-size: 18px; font-weight: bold; width: 100%; border: none; transition: 0.3s; }
    div.stButton > button:first-child:hover { background-color: #112d47; box-shadow: 0px 4px 10px rgba(0,0,0,0.2); }
    .stTextArea textarea { border-radius: 6px !important; border: 1px solid #999 !important; }
    </style>
    """, unsafe_allow_html=True)

# --- XML Tools ---
def set_cell_border(border_el, val='single', sz='12', space='0', color='000000'):
    border_el.set(qn('w:val'), val)
    border_el.set(qn('w:sz'), sz)
    border_el.set(qn('w:space'), space)
    border_el.set(qn('w:color'), color)

def add_continuous_section_break(paragraph, num_cols):
    pPr = paragraph._element.get_or_add_pPr()
    existing_sectPr = pPr.find(qn('w:sectPr'))
    if existing_sectPr is not None:
        pPr.remove(existing_sectPr)
        
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'continuous')
    sectPr.append(type_el)
    
    cols_el = OxmlElement('w:cols')
    cols_el.set(qn('w:num'), str(num_cols))
    cols_el.set(qn('w:space'), '400')
    if num_cols == 2:
        cols_el.set(qn('w:sep'), '1')
    sectPr.append(cols_el)
    
    margin = OxmlElement('w:pgMar')
    margin.set(qn('w:top'), '432')
    margin.set(qn('w:bottom'), '432')
    margin.set(qn('w:left'), '720')
    margin.set(qn('w:right'), '720')
    sectPr.append(margin)
    
    pPr.append(sectPr)

# --- હેડર ડિઝાઇન ---
def insert_header_table(doc, header_left, header_center):
    h_font = "Times New Roman"
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    
    first_para = doc.paragraphs[0]._p
    first_para.addprevious(table._tbl)
    
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.17) 
    table.columns[2].width = Inches(1.3)
    
    left_cell = table.cell(0, 0)
    p_logo_left = left_cell.paragraphs[0]
    p_logo_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists('logo.png'): p_logo_left.add_run().add_picture('logo.png', width=Inches(1.7))
        
    lines = header_left.strip().split('\n')
    if not lines: lines = ["JEE MAIN", "GUJARATI MEDIUM"]
    if len(lines) == 1: lines.append(" ")
    
    nested_table = left_cell.add_table(rows=2, cols=1)
    tblPr = nested_table._tbl.tblPr
    tblW = OxmlElement('w:tblW', {qn('w:w'): '5000', qn('w:type'): 'pct'}) 
    tblPr.append(tblW)
    
    tblBorders = OxmlElement('w:tblBorders')
    for b_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        b_el = OxmlElement(f'w:{b_name}', {qn('w:val'): 'single', qn('w:sz'): '12', qn('w:space'): '0', qn('w:color'): '000000'})
        tblBorders.append(b_el)
    tblPr.append(tblBorders)
    
    for r_idx in range(2):
        n_cell = nested_table.cell(r_idx, 0)
        tcPr = n_cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd', {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D9D9D9'})
        tcPr.append(shd)
        
        tcMar = OxmlElement('w:tcMar')
        for m in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{m}', {qn('w:w'): '0', qn('w:type'): 'dxa'})
            tcMar.append(el)
        tcPr.append(tcMar)
        
        np = n_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(0)
        np.paragraph_format.space_after = Pt(0)
        
        n_run = np.add_run(lines[r_idx])
        n_run.font.name = "Arial"
        n_run.font.bold = True
        n_run.font.size = Pt(14) if r_idx == 0 else Pt(10)
    
    center_cell = table.cell(0, 1)
    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_lines = header_center.strip().split('\n')
    for i, line in enumerate(c_lines):
        r = p_center.add_run(line)
        r.font.name = h_font
        r.font.bold = True
        if i == 0: r.font.size = Pt(18)       
        elif i == 1: r.font.size = Pt(16)     
        else: r.font.size = Pt(14)            
        if i < len(c_lines) - 1: p_center.add_run('\n')
            
    right_cell = table.cell(0, 2)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists('sblogo.png'): p_right.add_run().add_picture('sblogo.png', width=Inches(1.1))

    # હેડરની નીચેની બ્લેક લાઈન (અહીંથી આપણે સેક્શન બ્રેક કાઢી નાખ્યો છે જેથી પેજ ના છૂટે)
    p_line = doc.add_paragraph()
    table._tbl.addnext(p_line._p)
    pPr = p_line._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom', {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '000000'})
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    return p_line

# --- વર્ડ ફાઈલ ફોર્મેટિંગ (કોરા પેજનું બ્રહ્માસ્ત્ર ફિક્સ) ---
def set_formatting_and_margins(docx_filename, font_size, font_name, header_left, header_center):
    doc = Document(docx_filename)
    
    # ડિફોલ્ટ આખું પેપર 2-કોલમ સેટ કરો
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        sectPr = section._sectPr
        cols = sectPr.find(qn('w:cols')) or OxmlElement('w:cols')
        if cols not in sectPr: sectPr.append(cols)
        cols.set(qn('w:num'), '2')       
        cols.set(qn('w:space'), '400')   
        cols.set(qn('w:sep'), '1')       
    
    header_line_para = insert_header_table(doc, header_left, header_center)

    # નકામા ફકરા ડિલીટ કરો (હેડરની લાઈન સિવાય)
    for paragraph in list(doc.paragraphs):
        if paragraph._element == header_line_para._element:
            continue
        if not paragraph.text.strip():
            p = paragraph._element
            if p.getparent() is not None:
                p.getparent().remove(p)
            paragraph._p = paragraph._element = None
            
    # જીવંત ફકરાઓનું નવું લિસ્ટ
    paragraphs = [p for p in doc.paragraphs if p.text.strip() or p._element == header_line_para._element]
    
    is_one_col = True # શરૂઆતમાં હેડર 1-કોલમમાં હોય છે
    
    for i, paragraph in enumerate(paragraphs):
        if paragraph._element == header_line_para._element:
            continue
            
        # બ્રહ્માસ્ત્ર: ઓટો-બુલેટ કાયમ માટે બંધ
        paragraph.style = doc.styles['Normal']
        pPr = paragraph._element.get_or_add_pPr()
        numPrs = pPr.findall(qn('w:numPr'))
        for n in numPrs: pPr.remove(n)
            
        for run in paragraph.runs:
            if '‡' in run.text:
                run.text = run.text.replace('‡', '\t')
                
        text = paragraph.text.strip()
        
        # 1-Column ડાર્ક બ્લુ ટાઇટલ
        if '###HEADER###' in text:
            clean_title = text.replace('###HEADER###', '').strip()
            
            # જો અગાઉના પ્રશ્નો 2-કોલમમાં ચાલતા હતા, તો તેને ત્યાં પૂરા કરો
            if i > 0 and not is_one_col:
                prev_p = paragraphs[i-1]
                if prev_p._element == header_line_para._element and i > 1:
                    prev_p = paragraphs[i-2]
                add_continuous_section_break(prev_p, 2)
            
            # આ ટાઇટલને 1-કોલમમાં રાખો
            add_continuous_section_break(paragraph, 1)
            is_one_col = True
            
            paragraph.text = ""
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)
            
            shd = OxmlElement('w:shd', {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '1F4E79'})
            paragraph._element.get_or_add_pPr().append(shd)
            
            run = paragraph.add_run(clean_title)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(font_size + 4)
            run.font.name = "Times New Roman"
            continue

        else:
            # જો આ પહેલો જ પ્રશ્ન હોય (ટાઇટલ વગર), તો હેડરની લાઈન પછી 1-કોલમ બ્રેક મારો
            if is_one_col:
                add_continuous_section_break(header_line_para, 1)
                is_one_col = False

            # તમારું પોતાનું જ MCQ ફોર્મેટિંગ લોજિક
            if re.match(r'^Q\.\d+', text):
                paragraph.paragraph_format.left_indent = Inches(0.35)
                paragraph.paragraph_format.first_line_indent = Inches(-0.35)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(2) 
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.tab_stops.clear_all()
                paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.35), WD_TAB_ALIGNMENT.LEFT)
                
            elif re.match(r'^\(?[A-D][\)\.]', text):
                paragraph.paragraph_format.left_indent = Inches(0.35)
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                is_last_option = True
                for j in range(i + 1, len(paragraphs)):
                    next_text = paragraphs[j].text.strip()
                    if not next_text: continue
                    if re.match(r'^\(?[A-D][\)\.]', next_text):
                        is_last_option = False
                    break
                    
                paragraph.paragraph_format.space_after = Pt(8) if is_last_option else Pt(0)
                paragraph.paragraph_format.tab_stops.clear_all()
                tabs_count = paragraph.text.count('\t')
                
                if tabs_count == 3: 
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.8), WD_TAB_ALIGNMENT.LEFT)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(1.6), WD_TAB_ALIGNMENT.LEFT)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(2.4), WD_TAB_ALIGNMENT.LEFT)
                elif tabs_count == 1: 
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(1.6), WD_TAB_ALIGNMENT.LEFT)
            else:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                run.font.name = font_name
            
    # નવું ફૂટર અને વોટરમાર્ક લોજિક 
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists('sblogo.png'):
            try:
                image_part, rel_id = header.part.get_or_add_image('sblogo.png')
                watermark_xml = '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:pict><v:shape id="Watermark" style="position:absolute;left:0;text-align:left;margin-left:0;margin-top:0;width:350pt;height:350pt;z-index:-251657216;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" stroked="f"><v:imagedata r:id="' + rel_id + '" gain="35000f" blacklevel="15000f"/></v:shape></w:pict></w:r>'
                header_para._p.append(parse_xml(watermark_xml))
            except Exception: pass
            
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists('footer.png'):
            try:
                run = footer_para.add_run()
                run.add_picture('footer.png', width=Inches(7.5))
            except Exception: pass

    doc.save(docx_filename)

# --- 3. સ્માર્ટ માર્કડાઉન પાર્સર (તમારો જ અસલી કોડ 100% Copy-Paste) ---
def format_content(raw_text, is_continuous, start_num=1, end_num=0):
    raw_text = raw_text.replace('**', '')
    raw_text = re.sub(r'\n{3,}', '\n\n', raw_text)
    
    lines = raw_text.split('\n')
    questions = []
    current_q = []
    
    q_start_pattern = r'^[\s]*([Qq]\.?\s*\d+[\.\-\)]*|\d+[\.\-\)]+)\s+'
    
    for line in lines:
        if not line.strip(): continue
        
        if line.strip().startswith('#'):
            if current_q:
                questions.append("\n".join(current_q))
            questions.append(line.strip())
            current_q = []
        elif re.match(q_start_pattern, line):
            if current_q:
                questions.append("\n".join(current_q))
            current_q = [line]
        else:
            current_q.append(line)
            
    if current_q:
        questions.append("\n".join(current_q))
        
    formatted_md = ""
    q_num = start_num
    
    q_prefix_pattern = r'^([\s]*([Qq]\.?\s*\d+[\.\-\)]*|\d+[\.\-\)]+)\s*)+'
    labels = ['A', 'B', 'C', 'D']
    
    for q_block in questions:
        if end_num > 0 and q_num > end_num:
            q_num = 1
            
        if q_block.startswith('#'):
            if not is_continuous:
                q_num = start_num
                
            clean_title = re.sub(r'^#+', '', q_block).strip()
            formatted_md += f"###HEADER### {clean_title}\n\n"
            continue
            
        opt_pattern = r'\s*\(?[1-4A-Da-d][\)\.]\s*(.*?)(?=\s+\(?[1-4A-Da-d][\)\.]|$)'
        matches = list(re.finditer(opt_pattern, q_block, flags=re.DOTALL))
        
        if len(matches) >= 4:
            opts = matches[-4:]
            q_text = q_block[:opts[0].start()].strip()
            
            q_text = re.sub(q_prefix_pattern, '', q_text).strip()
            q_text = re.sub(r'\n\s*\n', '\n', q_text)
            
            q_md = f"**Q.{q_num}**‡{q_text}"
            q_num += 1
            
            clean_opts = []
            for i, m in enumerate(opts):
                opt_content = m.group(1).strip()
                opt_content = re.sub(r'\s+', ' ', opt_content) 
                clean_opts.append(f"\\({labels[i]}\\) {opt_content}")
                
            lens = [len(o) for o in clean_opts]
            max_len = max(lens)
            
            # તમારું 100% ઓરિજિનલ ગણિત
            if max_len < 16:
                opts_md = "‡".join(clean_opts)
            elif max_len < 36:
                opts_md = f"{clean_opts[0]}‡{clean_opts[1]}\n\n{clean_opts[2]}‡{clean_opts[3]}"
            else:
                opts_md = "\n\n".join(clean_opts)
                
            formatted_md += q_md + "\n\n" + opts_md + "\n\n"
        else:
            clean_q = re.sub(q_prefix_pattern, '', q_block).strip()
            if clean_q != q_block.strip() and re.match(q_start_pattern, q_block.strip()):
                 formatted_md += f"**Q.{q_num}**‡{clean_q}\n\n"
                 q_num += 1
            else:
                 formatted_md += q_block + "\n\n"
                 
    return formatted_md

# --- 4. Streamlit UI (Clean & Advanced) ---
missing = [f for f in ['logo.png', 'sblogo.png', 'footer.png'] if not os.path.exists(f)]
if missing: st.error(f"⚠️ ચેતવણી: આ ઈમેજ મિસિંગ છે: **{', '.join(missing)}**. પ્લીઝ અપલોડ કરો!")

col_logo1, col_logo2, col_logo3 = st.columns([2, 1, 2])
with col_logo2:
    if os.path.exists('logo.png'): st.image("logo.png", use_container_width=True) 

st.markdown("<h1 class='main-title'>Question Paper Generator</h1>", unsafe_allow_html=True)
st.markdown("<div class='subtitle-badge'><span>Made by Yug Ghanshyam Padmani</span></div>", unsafe_allow_html=True)

col_h1, col_h2 = st.columns(2)
with col_h1: header_left = st.text_area("૧. ડાબી બાજુનું ગ્રે બોક્સ (બે લાઈન):", "JEE MAIN\nGUJARATI MEDIUM", height=120)
with col_h2: header_center = st.text_area("૨. વચ્ચેનું ટાઈટલ (5 લાઈન):", "STD 11 SCIENCE\nMATHS\n40 MARKS\nJEE MAIN\nDate 13/08/26", height=120)

st.divider()

col1, col2 = st.columns([1, 2])
with col1:
    st.markdown("### ⚙️ ફાઇલ સેટિંગ્સ")
    file_name = st.text_input("ફાઈલનું નામ [ફરજિયાત]:", value="", placeholder="Physics_Test")
    font_name = st.selectbox("પેપરનો ફોન્ટ:", ["Hind Vadodara", "Shruti", "Times New Roman", "Arial"])
    font_size = st.number_input("ફોન્ટ સાઈઝ:", min_value=8, max_value=20, value=10)
    start_num = st.number_input("પ્રશ્ન ક્યાંથી શરૂ કરવો છે?", min_value=1, value=1)
    end_num = st.number_input("ક્યાં પૂરા કરવા છે? (0 એટલે બધા)", min_value=0, value=0)
    is_continuous = st.checkbox("પ્રશ્નોના નંબર સળંગ રાખવા છે?", value=True)

with col2:
    st.markdown("### ✍️ પ્રશ્નો પેસ્ટ કરો")
    user_input = st.text_area("પ્રશ્નો પેસ્ટ કરો (તમારું અસલી લોજિક):", height=350)

st.markdown("<br>", unsafe_allow_html=True)

if not shutil.which("libreoffice"):
    st.warning("⚠️ સિસ્ટમમાં PDF બનાવવાનું સોફ્ટવેર નથી.")
    if st.button("🔧 અત્યારે જ PDF સોફ્ટવેર ઇન્સ્ટોલ કરો (ફક્ત 1 મિનિટ લાગશે)"):
        with st.spinner("ઇન્સ્ટોલ થઈ રહ્યું છે... પ્લીઝ 1 મિનિટ રાહ જુઓ ⏳"):
            os.system("sudo apt-get update && sudo apt-get install libreoffice -y")
            st.success("✅ ઇન્સ્ટોલ થઈ ગયું! હવે તમે PDF બનાવી શકશો.")

if st.button("🚀 વર્ડ અને PDF ફાઇલ જનરેટ કરો"):
    if not file_name.strip(): st.error("⚠️ ભૂલ: 'ફાઈલનું નામ' ખાલી છે! સેટિંગ્સમાં જઈને નામ લખો.")
    elif not user_input.strip(): st.error("⚠️ ભૂલ: પ્રશ્નોનું બોક્સ ખાલી છે!")
    else:
        with st.spinner("તમારું પેપર બની રહ્યું છે... પ્લીઝ વેઇટ ⏳"):
            processed_md = format_content(user_input, is_continuous, start_num, end_num)
            with open("temp.md", "w", encoding="utf-8") as f: f.write(processed_md)
                
            try:
                subprocess.run(["pandoc", "temp.md", "-o", "temp.docx"], check=True)
                set_formatting_and_margins("temp.docx", font_size, font_name, header_left, header_center)
                final_file = f"{file_name}.docx"
                shutil.move("temp.docx", final_file)
                st.success("✅ તમારું પેપર સફળતાપૂર્વક બની ગયું છે!")
                
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    with open(final_file, "rb") as file:
                        st.download_button("📄 Word ફાઇલ ડાઉનલોડ કરો", file, file_name=final_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                with col_d2:
                    if shutil.which("libreoffice"):
                        try:
                            subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", final_file], check=True)
                            pdf_file = final_file.replace('.docx', '.pdf')
                            with open(pdf_file, "rb") as p_file:
                                st.download_button("📕 PDF ફાઇલ ડાઉનલોડ કરો", p_file, file_name=pdf_file, mime="application/pdf")
                        except Exception as e: st.error(f"⚠️ PDF કન્વર્ટ કરવામાં ભૂલ આવી: {e}")
                    else: st.error("⚠️ PDF સોફ્ટવેર નથી. પ્લીઝ ઉપર આપેલું 'ઇન્સ્ટોલ' બટન દબાવો.")
                        
            except Exception as e: st.error(f"Error: {e}")
